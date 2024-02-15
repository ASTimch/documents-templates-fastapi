from io import BytesIO
from typing import Any, Dict, Final, List, Literal, Tuple, TypeAlias

import docxtpl
import jinja2
import pymorphy2
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docxtpl import DocxTemplate
from num2words import num2words

morph = pymorphy2.MorphAnalyzer()


class InflectCase:
    """Падежи для фильтра."""

    nomn: Final[str] = "nomn"  # именительный
    gent: Final[str] = "gent"  # родительный
    datv: Final[str] = "datv"  # дательный
    accs: Final[str] = "accs"  # винительный
    ablt: Final[str] = "ablt"  # творительный
    loct: Final[str] = "loct"  # предложный
    voct: Final[str] = "voct"  # звательный


class CustomFilters:
    """Вспомогательные фильтры шаблонов."""

    def __init__(self):
        self._enabled = True
        self._skip_filter_tags = set()

    def enable(self, enable_filters: bool):
        """Активирует/деактивирует все пользовательские фильтры класса."""
        self._enabled = enable_filters

    def _skip_filter(self, tag: str) -> Tuple[bool, Any]:
        """Проверяет, нужно ли отключить/пропустить фильтр.

        Args:
            tag: Наименование тэга.

        Returns:
            (skip, value)
                skip (bool): True, если фильтр должен быть отключен.
                value (Any): значение возвращаемое вместо значения тега.
        """
        if (
            not self._enabled
            or tag is None
            or not tag
            or tag in self._skip_filter_tags
        ):
            return True, tag
        return False, tag

    def fio_short(self, fio: str) -> str:
        """Преобразует 'Фамилия Имя Отчество' к виду 'Фамилия И.О."""
        skip, value = self._skip_filter(fio)
        if skip:
            return value
        if not fio:
            return fio
        fields = fio.split(" ")
        initials = "".join(field[0].capitalize() + "." for field in fields[1:])
        return " ".join([fields[0].capitalize(), initials])

    def fio_title(self, fio: str) -> str:
        """Преобразует 'фамилия имя отчество' к виду 'Фамилия Имя Отчество"""
        skip, value = self._skip_filter(fio)
        if skip:
            return value
        if not fio:
            return fio
        return fio.title()

    def inflect_word(
        self,
        word: str,
        case: str,
    ) -> str:
        """Преобразование слова в заданный падеж.

        Args:
            word(str): Преобразуемое слово.
            case(str): Требуемый падеж.
                'nomn' именительный, 'gent' родительный, 'datv' дательный,
                'accs' винительный, 'ablt' творительный, 'loct' предложный,
                'voct' звательный.

        Returns:
            str: слово, преобразованное в заданный падеж.
        """
        if not word:
            return word
        try:
            p = next(
                filter(
                    lambda x: {InflectCase.nomn} in x.tag, morph.parse(word)
                )
            )
        except StopIteration:
            print("Not found 'nomn' form for ", word)
            return word
        return p.inflect({case}).word

    def inflect_words(self, words: str, case: str) -> str:
        """Преобразование каждого из слов в строке в заданный падеж.

        Args:
            words(str): Строка из нескольких слов.
            case(str): Требуемый падеж
                'nomn' именительный, 'gent' родительный, 'datv' дательный,
                'accs' винительный, 'ablt' творительный, 'loct' предложный,
                'voct' звательный.

        Returns:
            str: строка, каждое слово которого преобразовано в заданный падеж.
        """
        return " ".join(self.inflect_word(w, case) for w in words.split(" "))

    def genitive(self, words: str) -> str:
        """Преобразует все слова строки в родительный падеж."""
        skip, value = self._skip_filter(words)
        if skip:
            return value
        return self.inflect_words(words, InflectCase.gent)

    def dative(self, words: str) -> str:
        """Преобразует все слова строки в дательный падеж."""
        skip, value = self._skip_filter(words)
        if skip:
            return value
        return self.inflect_words(words, InflectCase.datv)

    def ablt(self, words: str) -> str:
        """Преобразует все слова строки в творительный падеж."""
        skip, value = self._skip_filter(words)
        if skip:
            return value
        return self.inflect_words(words, InflectCase.ablt)

    def loct(self, words: str) -> str:
        """Преобразует все слова строки в предложный падеж."""
        skip, value = self._skip_filter(words)
        if skip:
            return value
        return self.inflect_words(words, InflectCase.loct)

    def noun_plural(self, word: str, n: int) -> str:
        """Склонение заданного существительного в зависимости от числа n.

        Args:
            word (str): существительное.
            n (int): числительное.

        Returns:
            str: существительное склоненное в нужную форму для числительного n.
            Например: 'день' для n=1, 'дня' для n=2, 'дней' для n=5 и т.д.
        """
        skip, value = self._skip_filter(word)
        if skip:
            return value
        try:
            n = int(n)
        except Exception:
            return word
        word = morph.parse(word)[0]
        words = (
            word.inflect({"sing", InflectCase.nomn}).word,  # 'день'
            word.inflect({InflectCase.gent}).word,  # 'дня'
            word.inflect({"plur", InflectCase.gent}).word,  # 'дней'
        )

        n_mod100 = n % 100
        if n % 10 == 1 and n_mod100 != 11:
            return words[0]
        if 2 <= n % 10 <= 4 and (n_mod100 < 10 or n_mod100 >= 20):
            return words[1]
        return words[2]

    def adj_plural(self, word: str, n: int) -> str:
        """Склонение заданного прилагательного в зависимости от числа n.

        Args:
            word (str): существительное.
            n (int): числительное.

        Returns:
            str: прилагательное склоненное в нужную форму для числительного n.
            Например: 'новый' для n=[1, 21, 31..], 'новых' для n=2 и т.д.
        """
        skip, value = self._skip_filter(word)
        if skip:
            return value
        try:
            number = int(n)
        except Exception:
            return word
        word = morph.parse(word)[0]
        number_mod100 = number % 100
        if number % 10 == 1 and number_mod100 != 11:
            return word.inflect({"sing", InflectCase.nomn}).word  # 'новый'
        return word.inflect({"plur", InflectCase.gent}).word  # 'новых'

    def currency_to_words(self, num) -> str:
        """Преобразует заданную сумму в представление прописью."""
        if not self._enabled or num is None:
            return f"Прописью({num})"
        if num in self._skip_filter_tags:
            return f"Прописью({num})"
        try:
            roubles = round(float(num), 2)
        except Exception:
            return ""

        if roubles >= 10**14:
            return num
        return num2words(
            roubles, to="currency", lang="ru", cents=False, currency="RUB"
        )

    def split(self, line: str, sep=None) -> List[str]:
        """Разбивает строку на части и возвращает в виде списка слов."""
        skip, value = self._skip_filter(line)
        if skip:
            return [value]
        if not line or not isinstance(line, str):
            return line
        return line.split(sep)

    def get_filters(self):
        """Возвращает словарь вида {тег:функция} для всех фильтров."""
        return {
            "fio_short": self.fio_short,
            "fio_title": self.fio_title,
            "genitive": self.genitive,
            "dative": self.dative,
            "ablt": self.ablt,
            "loct": self.loct,
            "noun_plural": self.noun_plural,
            "adj_plural": self.adj_plural,
            "currency_to_words": self.currency_to_words,
            "split": self.split,
        }


class DocxRender:
    # Предопределенное наименование стиля для всех тэгов(переменных) в шаблоне
    TAG_STYLE_NAME: Final = "TemplateTag"

    def __init__(self, template_file_name: str):
        self._template_file_name = template_file_name
        self._template: DocxTemplate = docxtpl.DocxTemplate(template_file_name)
        self._jinja_env = jinja2.Environment()
        self._customfilters = CustomFilters()
        self._jinja_env.filters.update(self._customfilters.get_filters())

    def _render_to_file_stream(self, context: Dict[str, str]) -> BytesIO:
        """Генерирует docx документ из шаблона согласно контексту.

        Args:
            context: словарь вида {тэг:значение} для генерации документа

        Returns:
            BytesIO: документ после замены в шаблоне всех тегов на значения.
        """
        self._template.render(context, jinja_env=self._jinja_env)
        file_stream = BytesIO()
        self._template.save(file_stream)
        file_stream.seek(0)
        return file_stream

    def get_document(self, context: Dict[str, str]) -> BytesIO:
        """Генерирует и возвращает документ согласно заданному контексту.

        Args:
            context: словарь вида {тэг:значение} для генерации документа

        Returns:
            BytesIO: документ после замены в шаблоне всех тегов на значения.
        """
        self._customfilters.enable(True)  # switch custom filters on
        self._customfilters._skip_filter_tags.clear()
        return self._render_to_file_stream(context)

    def get_draft(self, context: Dict[str, str]) -> BytesIO:
        """Генерирует и возвращает эскиз документа согласно контексту

        Все поля тэгов маркируются желтым.

        Args:
            context: словарь вида {тэг:значение} для генерации эскиза.

        Returns:
            BytesIO: эскиз после замены в шаблоне всех тегов на значения.
        """
        self._customfilters.enable(False)  # switch custom filters off
        self.markdown_tags()
        # self._customfilters.enable(True)  # switch custom filters on
        return self._render_to_file_stream(context)

    def get_partial(
        self, context: Dict[str, str], context_default: Dict[str, str] = None
    ) -> BytesIO:
        """Генерирует и возвращает частично заполненный документ.

        Тэги, которые заданы в context_default и не найдены в context,
        маркируются желтым цветом (как не заполненные).

        Args:
            context: Словарь значения полей вида {тэг:значение}.
            context_default: Словарь значений по умолчанию вида
                {тэг: значение по умолчанию}.

        Returns:
            BytesIO: Документ после замены в шаблоне всех тегов на значения.
        """
        self._customfilters.enable(True)  # switch custom filters on
        self._template.init_docx()
        non_filled_tags = self.get_tags() - context.keys()
        default_tags = set()
        if context_default:
            default_tags = non_filled_tags & context_default.keys()
            self._markdown_given_tags(self._template.docx, default_tags)
            for tag in default_tags:
                self._customfilters._skip_filter_tags.add(context_default[tag])
                context[tag] = context_default[tag]
        self._customfilters.enable(True)
        return self._render_to_file_stream(context)

    def save(self, result_file_name):
        """Сохранение текущего документа."""
        self._template.save(result_file_name)

    def get_tags(self) -> List[str]:
        """Возвращает список всех тэгов из docx шаблона."""
        return self._template.get_undeclared_template_variables(
            jinja_env=self._jinja_env
        )

    def markdown_tags(self, color=WD_COLOR_INDEX.YELLOW):
        """Размечает места тэгов заданным цветом."""
        self._template.init_docx()
        self._markdown_tag(self._template.docx, color, "{{")
        self._markdown_tag(self._template.docx, color, "}}")
        # self._template.is_rendered = True

    def _docx_paragraphs(self, docx: Document):
        """Генератор по всем параграфам документа."""
        for p in docx.paragraphs:
            yield p
        for table in docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    def _docx_runs(self, docx: Document):
        """Генератор по всем прогонам (run) документа."""
        for p in self._docx_paragraphs(docx):
            for r in p.runs:
                yield r

    def _combine_styled_tag_runs(self, tag_style, runs):
        """Объединяет последовательные прогоны стиля tag_style в один."""
        start_run = None
        for r in runs:
            if r.style == tag_style:
                if start_run:
                    start_run.text += r.text
                    r.clear()
                else:
                    start_run = r
            else:
                start_run = None

    def prepare_template(self):
        """Подготовка шаблона к использованию (объединение прогонов)."""
        self._template.init_docx()
        docx = self._template.docx
        tag_style = docx.styles[self.TAG_STYLE_NAME]
        # self._print_document_runs(docx)
        runs = list(self._docx_runs(docx))
        # Объединение последовательных прогонов с тэгами в один
        self._combine_styled_tag_runs(tag_style, runs)
        docx.save(self._template_file_name)

    def _markdown_given_tags(
        self, docx: Document, tags: List[str], color=WD_COLOR_INDEX.YELLOW
    ):
        """Маркировка заданных тегов при помощи заданного цвета.

        Args:
            docx (Document): модифицируемый docx документ.
            tags (list[str]): список тэгов, который должны быть промаркированы.
            color: цвет для маркировки тэгов.
        """

        def markdown_tag_begin(i, runs):
            while i >= 0:
                if "{{" in runs[i].text:
                    runs[i].font.highlight_color = color
                    return
                i -= 1

        tags_set = set(tags)
        tag_style = docx.styles[self.TAG_STYLE_NAME]
        runs = list(self._docx_runs(docx))
        # TODO: подготовка должна быть выполнена при загрузке шаблона в базу
        self._combine_styled_tag_runs(tag_style, runs)
        for i, r in enumerate(runs):
            if r.style == tag_style and r.text in tags_set:
                markdown_tag_begin(i, runs)

    def _print_document_runs(self, docx: Document):
        """Анализ документа: печать всех его прогонов (run)"""
        invalid_runs = []
        for p in self._docx_paragraphs(docx):
            print(f"p=<{p.text}>")
            for r in p.runs:
                print(f"r=<{r.text}>")
                cnt_open = r.text.count("{")
                cnt_closed = r.text.count("}")
                if cnt_open % 2 or cnt_closed % 2:
                    invalid_runs.append(r.text)
        if invalid_runs:
            print("Invalid runs: ", invalid_runs)

    def _markdown_tag(self, docx: Document, color, tag: str = "{{"):
        """
        Подсветка заданного тега в документе при помощи заданного цвета.
        """
        for r in self._docx_runs(docx):
            if tag in r.text:
                r.font.highlight_color = color
